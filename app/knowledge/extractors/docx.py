from ..models import ExtractedDocument, ExtractedUnit, SourceMetadata
from .base import ExtractionError, ExtractionLimits, enforce_limits


def extract(path: str, source: SourceMetadata, limits: ExtractionLimits) -> ExtractedDocument:
    try:
        from docx import Document
        document = Document(path)
    except Exception as exc:
        raise ExtractionError(f"could not open DOCX: {exc}") from exc
    units, parts, heading, number = [], [], None, 1
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if parts:
                units.append(ExtractedUnit("section", str(number), "\n".join(parts), heading=heading))
                number += 1
                parts = []
            heading = text
        else:
            parts.append(text)
    for table in document.tables:
        parts.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
    if parts or heading:
        units.append(ExtractedUnit("section", str(number), "\n".join(parts) or heading or "", heading=heading))
    title = document.core_properties.title or source.display_name
    return enforce_limits(ExtractedDocument(title, "document", units), limits)
